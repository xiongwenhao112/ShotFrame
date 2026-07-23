# -*- coding: utf-8 -*-
"""测试素材自动生成：克隆仓库后无需任何外部文件即可跑全部测试。"""
import os

from PIL import Image

from shotframe.core import make_sample

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(HERE, "test_data")
IMG1 = os.path.join(DATA, "image1.png")
IMG9 = os.path.join(DATA, "image9.png")
DOCX = os.path.join(DATA, "测试文稿.docx")
DOCX_IMAGES = 3


def ensure_fixtures():
    """生成两张演示截图和一个内嵌 3 张图的 docx，已存在则跳过。"""
    os.makedirs(DATA, exist_ok=True)
    if not os.path.exists(IMG1):
        make_sample(940, 520).save(IMG1)
    if not os.path.exists(IMG9):
        make_sample(1030, 640).save(IMG9)
    if not os.path.exists(DOCX):
        make_docx(DOCX, DOCX_IMAGES)
    return IMG1, IMG9, DOCX


def make_docx(path, n_images=3):
    """用 python-docx 生成一个包含 n 张插图的测试文稿。"""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("ShotFrame 测试文稿", level=1)
    for i in range(n_images):
        doc.add_paragraph("下面是第 %d 张测试截图：" % (i + 1))
        img = os.path.join(DATA, "_docx_img%d.png" % i)
        make_sample(700 + i * 60, 420 + i * 30).save(img)
        doc.add_picture(img, width=Inches(5.5))
        doc.add_paragraph("这一段是截图后的说明文字。")
    doc.save(path)
    for i in range(n_images):
        img = os.path.join(DATA, "_docx_img%d.png" % i)
        if os.path.exists(img):
            os.remove(img)
    return path


MD = os.path.join(DATA, "测试文章.md")


def ensure_md_fixture():
    """生成引用一张本地图片的测试 Markdown。"""
    ensure_fixtures()
    if not os.path.exists(MD):
        with open(MD, "w", encoding="utf-8") as f:
            f.write("# 测试文章\n\n正文一段。\n\n"
                    "![示例](image1.png)\n\n又一段。\n")
    return MD


if __name__ == "__main__":
    print(ensure_fixtures())
    print(ensure_md_fixture())
